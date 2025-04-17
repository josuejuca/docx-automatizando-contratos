from docx import Document
from datetime import datetime
from num2words import num2words
import os
import uuid
import tempfile
import locale
import copy
from docx.oxml import OxmlElement
from docx.oxml.table import CT_Tbl

try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
except locale.Error:
    pass


from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


# NPS 

def inserir_avaliacao_nps(doc, avaliacoes: dict):
    criterios = [
        "Localização",
        "Tamanho",
        "Planta (disposição dos cômodos)",
        "Qualidade / Acabamentos",
        "Estado de Conservação",
        "Áreas comuns",
        "Preço"
    ]

    idx, tabela = encontrar_tabela_antes_placeholder(doc, "{avaliacao_nps}")
    if tabela is None:
        return

    linhas = tabela.rows

    for i, criterio in enumerate(criterios):
        nota = avaliacoes.get(criterio, 0)
        if nota >= 1 and nota <= 5:
            # linha i+1 porque a primeira linha é o cabeçalho
            celula = linhas[i + 1].cells[nota]  # coluna 1 a 5
            # Limpa o conteúdo anterior da célula
            celula.text = ""
            par = celula.paragraphs[0]
            run = par.add_run("X")
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    # Remove o parágrafo do placeholder
    for para in doc.paragraphs:
        if "{avaliacao_nps}" in para.text:
            para.text = ""
            break



def get_tabela_from_element(doc, element):
    for tbl in doc.tables:
        if isinstance(tbl._element, CT_Tbl) and tbl._element == element:
            return tbl
    return None

def encontrar_tabela_antes_placeholder(doc, placeholder):
    body = doc._body._element
    for i in range(len(body) - 1):
        el = body[i]
        next_el = body[i + 1]
        if next_el.tag.endswith("p") and next_el.text and placeholder in next_el.text:
            if el.tag.endswith("tbl"):
                num_tbls = len([e for e in body[:i + 1] if e.tag.endswith("tbl")])
                return i, doc.tables[num_tbls - 1]
    return None, None

def preencher_por_variaveis(tabela, contexto):
    for row in tabela.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for chave, valor in contexto.items():
                        if chave in run.text:
                            run.text = run.text.replace(chave, str(valor))

def substituir_texto_em_todo_documento(doc, chave, valor):
    for para in doc.paragraphs:
        if chave in para.text:
            para.text = para.text.replace(chave, str(valor))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if chave in para.text:
                        para.text = para.text.replace(chave, str(valor))

def inserir_enter_apos(elemento):
    p = OxmlElement("w:p")
    elemento.addnext(p)

def preencher_declaracao_visita(payload: dict, template_path: str) -> str:
    doc = Document(template_path)

    # Substituições dos dados do imóvel
    substituir_texto_em_todo_documento(doc, "endereco_imovel", payload["endereco_imovel"])    

    # Substituição da data por extenso
    hoje = datetime.today()
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    data_formatada = f"{hoje.day:02d} de {meses[hoje.month - 1]} de {hoje.year}"
    substituir_texto_em_todo_documento(doc, "{{data_full}}", data_formatada)

    # Gerar tabelas dos visitantes no local do {visitante_teble}
    idx, tabela_modelo = encontrar_tabela_antes_placeholder(doc, "{visitante_teble}")
    if idx is not None and tabela_modelo is not None:
        modelo = copy.deepcopy(tabela_modelo._element)
        doc._body._element.remove(tabela_modelo._element)

        for para in doc.paragraphs:
            if "{visitante_teble}" in para.text:
                ponto_insercao = OxmlElement("w:p")
                para._element.addnext(ponto_insercao)
                doc._body._element.remove(para._element)
                break

        for visitante in payload["visitantes"]:
            nova = copy.deepcopy(modelo)
            ponto_insercao.addprevious(nova)
            inserir_enter_apos(nova)
            tabela_nova = get_tabela_from_element(doc, nova)
            if tabela_nova:
                preencher_por_variaveis(tabela_nova, {
                    "nome_visitante": visitante["nome"],
                    "cpf_visitante": visitante["cpf"],
                    "email_visitante": visitante["email"],
                    "tel": visitante["tel"]
                })
    
    if payload.get("isImob"):
        texto_imob = f", (parceiro/associado) da imobiliária {payload['name_imob']}, inscrita no CRECI/DF sob o nº {payload['number_creci']}"
    else:
        texto_imob = ""

    substituir_texto_em_todo_documento(doc, "{imob}", texto_imob)

    # === Assinaturas dos visitantes ===
    for i, para in enumerate(doc.paragraphs):
        if "{assinatura_visitante}" in para.text:
            para.text = ""  # remove o placeholder
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for idx in range(0, len(payload["visitantes"]), 2):
                row = table.add_row().cells
                for j in range(2):
                    if idx + j < len(payload["visitantes"]):
                        visitante = payload["visitantes"][idx + j]
                        row[j].paragraphs[0].add_run("__________________________________")
                        row[j].add_paragraph(f"Nome: {visitante['nome']}")
                        row[j].add_paragraph(f"CPF: {visitante['cpf']}")
            para._element.addnext(table._tbl)
            break


    substituir_texto_em_todo_documento(doc, "NOME DO CORRETOR", payload["nome_corretor"])
    substituir_texto_em_todo_documento(doc, "CRECI DO CORRETOR(A)", payload["creci_corretor"])

    if "avaliacao_nps" in payload:
        inserir_avaliacao_nps(doc, payload["avaliacao_nps"])


    # Salvar
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, f"declaracao_visita_{uuid.uuid4()}.docx")
    doc.save(output_path)
    return output_path

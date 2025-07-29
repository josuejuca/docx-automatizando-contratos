from docx import Document
from typing import Dict, List, Optional
from io import BytesIO
import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
from docx.shared import Pt  # Importar isso no começo do arquivo!
from num2words import num2words
from datetime import datetime
import re
def encontrar_tabela_antes_placeholder(doc, placeholder: str):
    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            # Procura a última tabela antes do parágrafo com o placeholder
            for table in reversed(doc.tables):
                if table._element.getprevious() is para._element:
                    return table
            break
    return None

def preencher_pagamento_table(doc, pagamentos: List[Dict]):
    if not doc.tables:
        return

    tabela_pag = doc.tables[0]

    if tabela_pag:
        linha_modelo = tabela_pag.rows[1]
        for pagamento in pagamentos:
            nova_linha = tabela_pag.add_row()
            for i, cell in enumerate(nova_linha.cells):
                texto = linha_modelo.cells[i].text
                texto = texto.replace("tipo", pagamento["tipo"])
                texto = texto.replace("vencimento", pagamento["vencimento"])
                texto = texto.replace("valor", f'R$ {pagamento["valor"]:,.2f}'.replace(",", "X").replace(".", ",").replace("X", "."))
                texto = texto.replace("forma", pagamento["forma_pagamento"])
                texto = texto.replace("juros", pagamento.get("juros", "0%"))

                # Limpa o conteúdo atual e cria novo parágrafo com fonte 9
                cell.text = ""
                paragrafo = cell.paragraphs[0]
                run = paragrafo.add_run(texto)
                run.font.size = Pt(9)

        # Remove a linha modelo
        tabela_pag._tbl.remove(linha_modelo._tr)

    # Remove o parágrafo com {pagamento_table}
    for para in doc.paragraphs:
        if "{pagamento_table}" in para.text:
            doc._body._element.remove(para._element)
            break

def formatar_moeda(valor: float) -> str:
    if valor is None:
        return ""
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def valor_por_extenso(valor: float) -> str:

    if valor is None:
        return ""
    valor_int = int(valor)
    return num2words(valor_int, lang='pt-br').replace(" e zero", "") + " reais"

def gerar_texto_isimob(paragraph, is_imob: bool, nome_imob: Optional[str]):
    if "{{isImob}}" in paragraph.text:
        # print(f"[DEBUG] Parágrafo antes: {paragraph.text}")
        full_text = ''.join(run.text for run in paragraph.runs)
        partes = full_text.split("{{isImob}}")

        # Limpar o conteúdo atual do parágrafo
        for run in paragraph.runs:
            run.text = ""

        # Adicionar texto antes do placeholder
        paragraph.add_run(partes[0])

        # Adicionar substituição
        if is_imob and nome_imob:
            paragraph.add_run("da ")
            nome_run = paragraph.add_run(nome_imob)
            nome_run.bold = True
            paragraph.add_run(", e do(s) corretor(es) de imóveis parceiros/associados")
        else:
            paragraph.add_run("do(s) corretor(es) de imóveis")

        # Adicionar texto depois do placeholder
        if len(partes) > 1:
            paragraph.add_run(partes[1])

        # print(f"[DEBUG] Parágrafo depois: {''.join(run.text for run in paragraph.runs)}")



def substituir_variaveis_em_runs(paragraph, variaveis_bold: Dict[str, str]):
    full_text = ''.join(run.text for run in paragraph.runs)

    for placeholder, valor in variaveis_bold.items():
        if placeholder == "{gravame}" and placeholder in full_text:
            for idx, run in enumerate(paragraph.runs):
                if placeholder in run.text:
                    # Remove espaço no final do run anterior, se houver
                    if idx > 0 and paragraph.runs[idx-1].text.endswith(" "):
                        paragraph.runs[idx-1].text = paragraph.runs[idx-1].text.rstrip()
                    
                    parts = run.text.split(placeholder)
                    run.text = parts[0]
                    if len(parts) > 1:
                        paragraph.add_run(parts[1])

                    if valor:
                        paragraph.add_run(", com exceção da ")
                        paragraph.add_run(valor.get("tipo_gravame", ""))
                        paragraph.add_run(" em favor da(o) ")
                        paragraph.add_run(valor.get("beneficiario_gravame", ""))
                        paragraph.add_run(", inscrito no CNPJ/MF nº ")
                        paragraph.add_run(valor.get("beneficiario_cnpj_gravame", ""))
                        paragraph.add_run(", registrada no R-")
                        paragraph.add_run(valor.get("registro_gravame", ""))
                        paragraph.add_run(" da matrícula do imóvel no Cartório de Registro de Imóveis, cujo saldo devedor será quitado pelo banco financiador dessa transação (abatendo do valor a ser repassado ao(s) ")
                        run_bold1 = paragraph.add_run("PROMITENTE(S) VENDEDOR(ES)")
                        run_bold1.bold = True
                        paragraph.add_run(" pelo sistema de ")
                        run_bold2 = paragraph.add_run("Interveniente Quitante")
                        run_bold2.bold = True
                        paragraph.add_run(", e a consequente averbação da baixa da referida Alienação Fiduciária ocorrerá junto com o registro da presente Compra e Venda.")
                    return      
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, f"@@@{placeholder}@@@")

    if "@@@" not in full_text:
        return

    for run in paragraph.runs:
        run.text = ""

    partes = full_text.split("@@@")
    for parte in partes:
        if parte in variaveis_bold:
            valor_substituto = variaveis_bold[parte]
            if not isinstance(valor_substituto, str):
                valor_substituto = str(valor_substituto)
            paragraph.add_run(valor_substituto)
        else:
            paragraph.add_run(parte)

def montar_texto_pessoa(doc, lista_pessoas: List[Dict], paragrafo, tipo: str):
    total = len(lista_pessoas)
    for idx, p in enumerate(lista_pessoas):
        if idx > 0:
            if idx == total - 1:
                paragrafo.add_run(" e ")
            else:
                paragrafo.add_run(", ")

        run_nome = paragrafo.add_run(p['nome'])
        run_nome.bold = True
        paragrafo.add_run(f", {p['nacionalidade']}, portador(a) da carteira de identidade nº {p['rg_number']}, expedido pela {p['ssp_rg']}, inscrito(a) no CPF sob o nº {p['cpf']}, {p['estado_civil']}, ")

        if p.get('estado_civil', '').lower() in ["casado(a)", "casado(a)"] and all([
            p.get('nome_conjuge'),
            p.get('nacionalidade_conjuge'),
            p.get('rg_conjuge'),
            p.get('ssp_rg_conjuge'),
            p.get('cpf_conjuge')
        ]):
            paragrafo.add_run("com ")
            run_conjuge = paragrafo.add_run(p['nome_conjuge'])
            run_conjuge.bold = True
            paragrafo.add_run(f", {p['nacionalidade_conjuge']}, portador(a) da carteira de identidade nº {p['rg_conjuge']}, expedido pela {p['ssp_rg_conjuge']}, inscrito(a) no CPF sob o nº {p['cpf_conjuge']}, ")

        paragrafo.add_run(f"residente(s) e domiciliado(a)(s) no endereço {p['endereco']}, telefone {p['telefone']}" )

    paragrafo.add_run(", doravante denominado(a)(s) simplesmente, ")
    if tipo == "vendedores":
        run_final = paragrafo.add_run("PROMITENTE(S) VENDEDOR(A)(ES).")
    else:
        run_final = paragrafo.add_run("PROMISSÁRIO(A)(S) COMPRADOR(A)(ES).")
    run_final.bold = True

def gerar_texto_fgts(paragraph, fgts_ativo: bool):
    """Substitui {fgts} no parágrafo, sem misturar com outras funções."""

    if "{fgts}" in paragraph.text:
        # Apaga todos os runs atuais
        for run in paragraph.runs:
            run.text = ""

        if fgts_ativo:
            # FGTS ativo = insere o texto completo
            paragraph.add_run("2.7. ")
            run1 = paragraph.add_run("O(s) PROMISSÁRIO(S) COMPRADOR(ES) ")
            run1.bold = True
            paragraph.add_run("se responsabiliza(m) plenamente pela concessão do financiamento e/ou saque do FGTS, que poderá ser obtido junto ao Agente Financiador que escolher(em), obrigando-se o(s) PROMISSÁRIO(S) COMPRADOR(ES) a prover(em) o pagamento da referida parcela, ou de parte dela, caso não financie(m) o valor total, com recursos próprios.\n\n")
            run2 = paragraph.add_run("Parágrafo Primeiro: ")
            run2.bold = True
            paragraph.add_run("o pagamento de eventuais tarifas e despesas decorrentes do processo de financiamento serão de inteira responsabilidade do(s) PROMISSÁRIO(S) COMPRADOR(ES).")
        else:
            # FGTS falso = só limpa (remove o placeholder)
            paragraph.add_run("")

def gerar_texto_relacao_movies(paragraph, texto_relacao: Optional[str]):
    """Substitui {relacao_movies} no parágrafo, se preenchido."""

    if "{relacao_movies}" in paragraph.text:
        # Apaga o conteúdo atual
        for run in paragraph.runs:
            run.text = ""

        if texto_relacao:
            paragraph.add_run(f'"{texto_relacao}"')  # Coloca as aspas automáticas
        else:
            paragraph.add_run("")  # Deixa vazio se o usuário não preencher


def preencher_promessa(dados_vendedores: List[Dict], dados_compradores: List[Dict], dados_imovel: Dict, pagamentos: List[Dict], modelo_path: str, dados_testemunhas: List[Dict]) -> BytesIO:

    doc = Document(modelo_path)

    gravame_info = None
    if dados_imovel.get("gravame"):
        gravame_info = {
            "tipo_gravame": dados_imovel.get("tipo_gravame", ""),
            "beneficiario_gravame": dados_imovel.get("beneficiario_gravame", ""),
            "beneficiario_cnpj_gravame": dados_imovel.get("beneficiario_cnpj_gravame", ""),
            "registro_gravame": dados_imovel.get("registro_gravame", "")
        }

    variaveis_bold = {
        "{endereco_imovel}": dados_imovel.get("endereco_imovel", ""),
        "{matricula_imovel}": dados_imovel.get("matricula_imovel", ""),
        "{numero_cartorio}": dados_imovel.get("numero_cartorio", ""),
        "{forma_de_pagamento_sinal}": dados_imovel.get("forma_de_pagamento_sinal", ""),
        "{gravame}": gravame_info,
        "{valor_imovel}": formatar_moeda(dados_imovel.get("valor_imovel")),
        "{valor_imovel_texto}": valor_por_extenso(dados_imovel.get("valor_imovel")),
        "{valor_sinal}": formatar_moeda(dados_imovel.get("valor_sinal")),
        "{valor_sinal_texto}": valor_por_extenso(dados_imovel.get("valor_sinal")),
        "{valor_comissao}": formatar_moeda(dados_imovel.get("valor_comissao")),
        "{valor_comissao_texto}": valor_por_extenso(dados_imovel.get("valor_comissao")),
        "{data_full}": datetime.now().strftime("%d de %B de %Y").lower()
    }

    for p in doc.paragraphs:
        if "{vendedores}" in p.text:
            p.text = ""
            montar_texto_pessoa(doc, dados_vendedores, p, tipo="vendedores")
        elif "{compradores}" in p.text:
            p.text = ""
            montar_texto_pessoa(doc, dados_compradores, p, tipo="compradores")
        else:
            substituir_variaveis_em_runs(p, variaveis_bold)
            gerar_texto_fgts(p, dados_imovel.get("fgts", False))
            gerar_texto_relacao_movies(p, dados_imovel.get("relacao_movies", ""))
            gerar_texto_isimob(p, dados_imovel.get("isImob", False), dados_imovel.get("nomeImob", ""))


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{vendedores}" in para.text:
                        para.text = ""
                        montar_texto_pessoa(doc, dados_vendedores, para, tipo="vendedores")
                    elif "{compradores}" in para.text:
                        para.text = ""
                        montar_texto_pessoa(doc, dados_compradores, para, tipo="compradores")
                    else:                        
                        substituir_variaveis_em_runs(para, variaveis_bold)
                        gerar_texto_fgts(para, dados_imovel.get("fgts", False))
                        gerar_texto_relacao_movies(para, dados_imovel.get("relacao_movies", ""))
                        gerar_texto_isimob(para, dados_imovel.get("isImob", False), dados_imovel.get("nomeImob", ""))

    preencher_pagamento_table(doc, pagamentos)

     # === ASSINATURA DOS VENDEDORES ===
    for para in doc.paragraphs:
        if "{assinatura_vendedor}" in para.text:
            para.text = ""
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for idx in range(0, len(dados_vendedores), 2):
                row = table.add_row().cells
                for j in range(2):
                    if idx + j < len(dados_vendedores):
                        vendedor = dados_vendedores[idx + j]
                        row[j].paragraphs[0].add_run("__________________________________")
                        row[j].add_paragraph(f"Nome: {vendedor['nome']}")
                        row[j].add_paragraph(f"CPF: {vendedor['cpf']}")
            para._element.addnext(table._tbl)
            break

    # === ASSINATURA DOS COMPRADORES ===
    for para in doc.paragraphs:
        if "{assinatura_comprador}" in para.text:
            para.text = ""
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for idx in range(0, len(dados_compradores), 2):
                row = table.add_row().cells
                for j in range(2):
                    if idx + j < len(dados_compradores):
                        comprador = dados_compradores[idx + j]
                        row[j].paragraphs[0].add_run("__________________________________")
                        row[j].add_paragraph(f"Nome: {comprador['nome']}")
                        row[j].add_paragraph(f"CPF: {comprador['cpf']}")
            para._element.addnext(table._tbl)
            break
    # === TESTEMUNHAS ===
    if dados_testemunhas and len(dados_testemunhas) >= 2:
        def substituir_texto_em_todo_documento(doc, antigo, novo):
            for p in doc.paragraphs:
                if antigo in p.text:
                    for run in p.runs:
                        if antigo in run.text:
                            run.text = run.text.replace(antigo, novo)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if antigo in p.text:
                                for run in p.runs:
                                    if antigo in run.text:
                                        run.text = run.text.replace(antigo, novo)

        substituir_texto_em_todo_documento(doc, "testemunha_1", dados_testemunhas[0]["nome"])
        substituir_texto_em_todo_documento(doc, "cpf_1_testemunha", dados_testemunhas[0]["cpf"])

        substituir_texto_em_todo_documento(doc, "testemunha_2", dados_testemunhas[1]["nome"])
        substituir_texto_em_todo_documento(doc, "cpf_2_testemunha", dados_testemunhas[1]["cpf"])    
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

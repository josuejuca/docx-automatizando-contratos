# preencher_contrato.py
from docx import Document
from docx.oxml import OxmlElement
import copy
from io import BytesIO
import re
import unicodedata
from typing import List
from docx.table import Table
from docx.oxml.table import CT_Tbl
from datetime import datetime
from num2words import num2words  # certifique-se que está instalado localmente
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def valor_extenso(valor: float) -> str:
    return num2words(valor, lang='pt_BR', to='currency')

def get_tabela_from_element(doc, element):
    for tbl in doc.tables:
        if isinstance(tbl._element, CT_Tbl) and tbl._element == element:
            return tbl
    return None

def normalizar(texto):
    return unicodedata.normalize('NFKD', texto).encode('ASCII', 'ignore').decode().lower()

def substituir_variaveis_em_runs(doc, contexto: dict):
    def substituir_em_paragrafos(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                for chave, valor in contexto.items():
                    if f"{chave}" in run.text:
                        run.text = run.text.replace(f"{chave}", valor)

    substituir_em_paragrafos(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_em_paragrafos(cell.paragraphs)

def preencher_por_variaveis(tabela, contexto):
    for row in tabela.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for chave, valor in contexto.items():
                        if chave in run.text:
                            run.text = run.text.replace(chave, str(valor))

def substituir_texto_em_todo_documento(doc, placeholder, novo_valor):
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, str(novo_valor))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(novo_valor))

def substituir_endereco(doc, endereco):
    for para in doc.paragraphs:
        if "{endereco_imovel}" in para.text:
            para.text = para.text.replace("{endereco_imovel}", endereco)

def encontrar_tabela_antes_placeholder(doc, placeholder):
    body = doc._body._element
    for i in range(len(body) - 1):
        el = body[i]
        next_el = body[i + 1]
        if next_el.tag.endswith("p") and next_el.text and placeholder in next_el.text:
            if el.tag.endswith("tbl"):
                num_tbls = len([e for e in body[:i + 1] if e.tag.endswith("tbl")])
                return i, doc.tables[num_tbls - 1]
    return None, None

def inserir_enter_apos(elemento):
    p = OxmlElement("w:p")
    elemento.addnext(p)

def preencher_por_rotulo(tabela, dados):
    for row in tabela.rows:
        for cell in row.cells:
            texto_cell = normalizar(cell.text)
            for chave, valor in dados.items():
                if normalizar(f"{chave}:") in texto_cell:
                    cell.text = f"{chave}: {valor}"

def preencher_contrato(
    endereco_imovel: str,
    contratantes: List[dict],
    corretores: List[dict],
    valor_venda: float,
    porcentagem_corretagem: float,
    testemunhas: List[dict],
    modelo_path="templates/contrato-de-corretagem.docx"
):
    doc = Document(modelo_path)

    substituir_texto_em_todo_documento(doc, "{endereco_imovel}", endereco_imovel)

    idx_c, tabela_c = encontrar_tabela_antes_placeholder(doc, "{contratante_table}")
    if idx_c is not None and tabela_c is not None:
        modelo_copia = copy.deepcopy(tabela_c._element)
        doc._body._element.remove(tabela_c._element)
        for para in doc.paragraphs:
            if "{contratante_table}" in para.text:
                ponto_insercao = OxmlElement("w:p")
                para._element.addnext(ponto_insercao)
                doc._body._element.remove(para._element)
                break
        for c in contratantes:
            nova = copy.deepcopy(modelo_copia)
            ponto_insercao.addprevious(nova)
            inserir_enter_apos(nova)
            tabela_inserida = get_tabela_from_element(doc, nova)
            if tabela_inserida:
                preencher_por_variaveis(tabela_inserida, {
                    "nome_prop": c['nome'],
                    "email_prop": c['email'],
                    "endereco_prop": c['endereco'],
                    "cpf_prop": c['cpf'],
                    "tel_prop": c['telefone'],
                    "cidade_prop": c['cidade'],
                    "cep_prop": c['cep'],
                    "uf_prop": c['uf'],
                })

    idx_k, tabela_k = encontrar_tabela_antes_placeholder(doc, "{contratado_table}")
    if idx_k is not None and tabela_k is not None:
        modelo_copia = copy.deepcopy(tabela_k._element)
        doc._body._element.remove(tabela_k._element)
        for para in doc.paragraphs:
            if "{contratado_table}" in para.text:
                ponto_insercao = OxmlElement("w:p")
                para._element.addnext(ponto_insercao)
                doc._body._element.remove(para._element)
                break
        for c in corretores:
            nova = copy.deepcopy(modelo_copia)
            ponto_insercao.addprevious(nova)
            inserir_enter_apos(nova)
            tabela_inserida = get_tabela_from_element(doc, nova)
            if tabela_inserida:
                preencher_por_variaveis(tabela_inserida, {
                    "nome_corretor": c['nome'],
                    "cnpj_corretor": c['cnpj'],
                    "endereco_corretor": c['endereco'],
                    "tel_corretor": c['telefone'],
                    "creci_corretor": c['creci'],
                })

    substituir_variaveis_em_runs(doc, {
        "nome_prop": contratantes[0]['nome'],
        "email_prop": contratantes[0]['email'],
        "endereco_prop": contratantes[0]['endereco'],
        "cpf_prop": contratantes[0]['cpf'],
        "tel_prop": contratantes[0]['telefone'],
        "cidade_prop": contratantes[0]['cidade'],
        "cep_prop": contratantes[0]['cep'],
        "uf_prop": contratantes[0]['uf'],
        "nome_corretor": corretores[0]['nome'],
        "cnpj_corretor": corretores[0]['cnpj'],
        "endereco_corretor": corretores[0]['endereco'],
        "tel_corretor": corretores[0]['telefone'],
        "creci_corretor": corretores[0]['creci'],
    })

    valor_comissao = valor_venda * (porcentagem_corretagem / 100)
    valor_comissao_texto = valor_extenso(valor_comissao)

    substituir_texto_em_todo_documento(doc, "{valor_comissao}", f"R$ {valor_comissao:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
    substituir_texto_em_todo_documento(doc, "{valor_comissao_texto}", valor_comissao_texto)

    idx_com, tabela_com = encontrar_tabela_antes_placeholder(doc, "{comissao_table}")
    if idx_com is not None and tabela_com is not None:
        linha_modelo = tabela_com.rows[1]
        for c in corretores:
            nova_linha = tabela_com.add_row()
            for i, cell in enumerate(nova_linha.cells):
                texto = linha_modelo.cells[i].text
                texto = texto.replace("name_corretor", c["nome"])
                texto = texto.replace("cnpj_number_comissao", c["cnpj"])
                texto = texto.replace("corretagem_percentual", f'{c["participacao"]:.0f}%')
                valor_individual = valor_comissao * (c["participacao"] / 100)
                texto = texto.replace("valor_corretagem", f'R$ {valor_individual:,.2f}'.replace(",", "X").replace(".", ",").replace("X", "."))
                cell.text = texto
        tabela_com._tbl.remove(linha_modelo._tr)
        for para in doc.paragraphs:
            if "{comissao_table}" in para.text:
                doc._body._element.remove(para._element)
                break

    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    hoje = datetime.today()
    data_formatada = f"{hoje.day:02d} de {meses[hoje.month - 1]} de {hoje.year}"
    substituir_texto_em_todo_documento(doc, "{data}", data_formatada)

    # ASSINATURA DOS PROPRIETÁRIOS
    for i, para in enumerate(doc.paragraphs):
        if "{assinatura_prop}" in para.text:
            para.text = ""  # remove o placeholder
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for idx in range(0, len(contratantes), 2):
                row = table.add_row().cells
                for j in range(2):
                    if idx + j < len(contratantes):
                        c = contratantes[idx + j]
                        row[j].paragraphs[0].add_run("__________________________________")
                        row[j].add_paragraph(f"Nome: {c['nome']}")
                        row[j].add_paragraph(f"CPF: {c['cpf']}")
            para._element.addnext(table._tbl)
            break
    # ASSINATURA DOS CORRETORES
    for i, para in enumerate(doc.paragraphs):
        if "{assinatura_corretor}" in para.text:
            para.text = ""  # remove o placeholder
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for idx in range(0, len(corretores), 2):
                row = table.add_row().cells
                for j in range(2):
                    if idx + j < len(corretores):
                        c = corretores[idx + j]
                        row[j].paragraphs[0].add_run("__________________________________")
                        row[j].add_paragraph(f"Nome: {c['nome']}")
                        row[j].add_paragraph(f"CNPJ: {c['cnpj']}")
            para._element.addnext(table._tbl)
            break   

    # === Substituir dados das testemunhas ===
    if len(testemunhas) >= 2:
        substituir_texto_em_todo_documento(doc, "testemunha_1", testemunhas[0].nome)
        substituir_texto_em_todo_documento(doc, "rg_1_testemunha", testemunhas[0].rg)
        substituir_texto_em_todo_documento(doc, "cpf_1_testemunha", testemunhas[0].cpf)

        substituir_texto_em_todo_documento(doc, "testemunha_2", testemunhas[1].nome)
        substituir_texto_em_todo_documento(doc, "rg_2_testemunha", testemunhas[1].rg)
        substituir_texto_em_todo_documento(doc, "cpf_2_testemunha", testemunhas[1].cpf)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
